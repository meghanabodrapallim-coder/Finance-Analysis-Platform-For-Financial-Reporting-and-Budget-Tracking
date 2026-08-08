from flask import send_file
import csv
import io
import os
import calendar
from openpyxl import Workbook
from docx import Document
from docx.shared import Pt, RGBColor
from flask import Flask, render_template, request, redirect, url_for, session
import sqlite3
from datetime import datetime, timedelta
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from werkzeug.security import generate_password_hash, check_password_hash

app = Flask(__name__)
app.jinja_env.globals.update(enumerate=enumerate)
app.secret_key = os.environ.get('SECRET_KEY', 'smartfinance123')
app.config['SESSION_COOKIE_HTTPONLY'] = True
app.config['SESSION_COOKIE_SAMESITE'] = 'Lax'
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(hours=2)
app.config['TESTING'] = False
DB_NAME = 'finance.db'

def init_db():
    conn = sqlite3.connect('finance.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                (id INTEGER PRIMARY KEY,
                name TEXT,
                email TEXT UNIQUE,
                password TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS expenses
                (id INTEGER PRIMARY KEY,
                user_id INTEGER,
                category TEXT,
                amount REAL,
                date TEXT,
                description TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS budget
                (id INTEGER PRIMARY KEY,
                user_id INTEGER,
                category TEXT,
                amount REAL,
                month TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS income
                (id INTEGER PRIMARY KEY,
                user_id INTEGER,
                amount REAL,
                month TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS investments
                (id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                asset TEXT,
                investment_name TEXT,
                invested_amount REAL,
                current_value REAL)''')
    c.execute('''CREATE TABLE IF NOT EXISTS goals
                (id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                goal_name TEXT,
                target_amount REAL,
                saved_amount REAL,
                target_date TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS notifications
                (id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER,
                type TEXT,
                message TEXT,
                priority TEXT,
                status TEXT,
                created_date TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS settings
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER UNIQUE,
            currency TEXT DEFAULT "INR",
            default_income REAL DEFAULT 0,
            notification_pref TEXT DEFAULT "on",
            theme TEXT DEFAULT "light")''')
    c.execute('''CREATE TABLE IF NOT EXISTS wallets
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            wallet_name TEXT,
            wallet_type TEXT,
            balance REAL DEFAULT 0)''')
    c.execute('''CREATE TABLE IF NOT EXISTS bills
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            bill_name TEXT,
            amount REAL,
            due_date TEXT,
            recurring TEXT DEFAULT 'No',
            paid TEXT DEFAULT 'No')''')
    c.execute('''CREATE TABLE IF NOT EXISTS subscriptions
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            subscription_name TEXT,
            amount REAL,
            billing_cycle TEXT,
            next_billing_date TEXT,
            category TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS challenges
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            challenge_name TEXT,
            challenge_type TEXT,
            target_days INTEGER,
            start_date TEXT,
            status TEXT DEFAULT 'Active')''')
    c.execute('''CREATE TABLE IF NOT EXISTS challenge_logs
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            challenge_id INTEGER,
            log_date TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS challenge_logs
            (id INTEGER PRIMARY KEY AUTOINCREMENT,
            challenge_id INTEGER,
            log_date TEXT)''')
    try:
        c.execute("ALTER TABLE settings ADD COLUMN notification_pref TEXT DEFAULT 'on'")
    except sqlite3.OperationalError:
        pass
    try:
        c.execute("ALTER TABLE settings ADD COLUMN theme TEXT DEFAULT 'light'")
    except sqlite3.OperationalError:
        pass
    # Performance indexes
    c.execute("CREATE INDEX IF NOT EXISTS idx_expenses_user ON expenses(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_expenses_date ON expenses(user_id, date)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_budget_user ON budget(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_investments_user ON investments(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_goals_user ON goals(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_income_user ON income(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_bills_user ON bills(user_id)")
    c.execute("CREATE INDEX IF NOT EXISTS idx_subscriptions_user ON subscriptions(user_id)")
    conn.commit()
    conn.close()

def get_db():
    conn = sqlite3.connect('finance.db')
    conn.row_factory = sqlite3.Row
    return conn

def get_full_report_data(user_id):
    conn = get_db()
    expenses = conn.execute(
        "SELECT * FROM expenses WHERE user_id=? ORDER BY date DESC", (user_id,)).fetchall()
    budgets = conn.execute(
        "SELECT * FROM budget WHERE user_id=?", (user_id,)).fetchall()
    investments = conn.execute(
        "SELECT * FROM investments WHERE user_id=?", (user_id,)).fetchall()
    goals = conn.execute(
        "SELECT * FROM goals WHERE user_id=?", (user_id,)).fetchall()
    conn.close()
    return expenses, budgets, investments, goals

@app.route('/')
def home():
    return redirect(url_for('login'))
@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        name = request.form['name'].strip()
        email = request.form['email'].strip().lower()
        raw_password = request.form['password']

        if not name or not email or not raw_password:
            return render_template('register.html', error="All fields are required.")
        if '@' not in email or '.' not in email:
            return render_template('register.html', error="Enter a valid email address.")
        if len(raw_password) < 4:
            return render_template('register.html', error="Password must be at least 4 characters.")

        password = generate_password_hash(raw_password)
        conn = get_db()
        try:
            conn.execute(
                "INSERT INTO users (name, email, password) VALUES (?, ?, ?)",
                (name, email, password))
            conn.commit()
            return redirect(url_for('login'))
        except:
            return render_template('register.html',
                                   error="Email already exists!")
        finally:
            conn.close()
    return render_template('register.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        email = request.form['email']
        password = request.form['password']
        conn = get_db()
        user = conn.execute(
            "SELECT * FROM users WHERE email=?",
            (email,)).fetchone()
        conn.close()
        if user and check_password_hash(user['password'], password):
            session.permanent = True
            session['user_id'] = user['id']
            session['user_name'] = user['name']
            return redirect(url_for('dashboard'))
    else:
            return render_template('login.html',
                                   error="Invalid email or password!")
    return render_template('login.html')

@app.route('/dashboard')
def dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user_id = session['user_id']
    conn = get_db()

    # Greeting
    current_hour = datetime.now().hour
    if current_hour < 12:
        greeting = "Good Morning"
    elif current_hour < 17:
        greeting = "Good Afternoon"
    else:
        greeting = "Good Evening"

    # Search filters
    search_category = request.args.get('search_category', '')
    search_date_from = request.args.get('search_date_from', '')
    search_date_to = request.args.get('search_date_to', '')

    query = "SELECT * FROM expenses WHERE user_id=?"
    params = [user_id]

    if search_category:
        query += " AND category=?"
        params.append(search_category)

    if search_date_from:
        query += " AND date>=?"
        params.append(search_date_from)

    if search_date_to:
        query += " AND date<=?"
        params.append(search_date_to)

    query += " ORDER BY date DESC"

    if not (search_category or search_date_from or search_date_to):
        query += " LIMIT 10"

    expenses = conn.execute(query, params).fetchall()

    # Total expenses
    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)
    ).fetchone()[0] or 0

    total_exp = float(total_exp)

    # Latest monthly income
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)
    ).fetchone()

    monthly_income = float(income_row["amount"]) if income_row else 0

    # Category-wise expenses
    cat_expenses = conn.execute(
        "SELECT category, SUM(amount) AS total FROM expenses WHERE user_id=? GROUP BY category",
        (user_id,)
    ).fetchall()

    # Calculations
    savings = monthly_income - total_exp

    budget_used = (
        round((total_exp / monthly_income) * 100, 1)
        if monthly_income > 0 else 0
    )

    health_score = max(0, min(100, int(100 - budget_used)))

    cat_labels = [row["category"] for row in cat_expenses]
    cat_amounts = [float(row["total"]) for row in cat_expenses]

    # Notifications
    notif_count = conn.execute(
        "SELECT COUNT(*) FROM notifications WHERE user_id=? AND status='Active'",
        (user_id,)
    ).fetchone()[0] or 0

    notifications = conn.execute(
        "SELECT * FROM notifications WHERE user_id=? AND status='Active' ORDER BY id DESC",
        (user_id,)
    ).fetchall()

    # Bill reminders
    reminder_count = conn.execute(
        "SELECT COUNT(*) FROM bills WHERE user_id=? AND paid='No' AND due_date <= ?",
        (
            user_id,
            (datetime.now() + timedelta(days=7)).strftime('%Y-%m-%d')
        )
    ).fetchone()[0]

    conn.close()

    return render_template(
        'dashboard.html',
        greeting=greeting,
        expenses=expenses,
        total_expenses=total_exp,
        monthly_income=monthly_income,
        savings=savings,
        budget_used=budget_used,
        health_score=health_score,
        cat_labels=cat_labels,
        cat_amounts=cat_amounts,
        notif_count=notif_count,
        reminder_count=reminder_count,
        notifications=notifications,
        search_category=search_category,
        search_date_from=search_date_from,
        search_date_to=search_date_to
    )
@app.route('/set_income', methods=['POST'])
def set_income():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    amount = float(request.form['income'])
    conn = get_db()
    conn.execute(
        "DELETE FROM income WHERE user_id=?", (user_id,))
    conn.execute(
        "INSERT INTO income (user_id, amount, month) VALUES (?, ?, '2026-07')",
        (user_id, amount))
    conn.commit()
    conn.close()
    return redirect(url_for('dashboard'))

@app.route('/add_expense', methods=['GET', 'POST'])
def add_expense():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if request.method == 'POST':
        user_id = session['user_id']
        category = request.form['category'].strip()
        try:
            amount = float(request.form['amount'])
        except ValueError:
            return render_template('expenses.html', error="Amount must be a valid number.")

        if amount <= 0:
            return render_template('expenses.html', error="Amount must be greater than zero.")
        if not category:
            return render_template('expenses.html', error="Category is required.")

        date = request.form['date']
        description = request.form['description'].strip()
        conn = get_db()
        conn.execute(
            "INSERT INTO expenses (user_id, category, amount, date, description) VALUES (?, ?, ?, ?, ?)",
            (user_id, category, amount, date, description))
        conn.commit()
        conn.close()
        return redirect(url_for('dashboard'))
    return render_template('expenses.html')

@app.route('/budget', methods=['GET', 'POST'])
def budget():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    if request.method == 'POST':
        category = request.form['category']
        amount = float(request.form['amount'])
        month = request.form['month']
        conn.execute(
            "INSERT INTO budget (user_id, category, amount, month) VALUES (?, ?, ?, ?)",
            (user_id, category, amount, month))
        conn.commit()

    budgets = conn.execute(
        "SELECT * FROM budget WHERE user_id=?",
        (user_id,)).fetchall()

    budget_data = []
    for b in budgets:
        spent = conn.execute(
            "SELECT SUM(amount) FROM expenses WHERE user_id=? AND category=?",
            (user_id, b['category'])).fetchone()[0] or 0
        spent = float(spent)
        remaining = float(b['amount']) - spent
        status = "On Track" if remaining >= 0 else "Over Budget"
        budget_data.append({
            'id': b['id'],
            'category': b['category'],
            'budget': float(b['amount']),
            'spent': spent,
            'remaining': remaining,
            'status': status
        })
    conn.close()
    return render_template('budget.html', budget_data=budget_data)

@app.route('/ai_analysis')
def ai_analysis():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row else 0

    cat_expenses = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category ORDER BY total DESC",
        (user_id,)).fetchall()
    conn.close()

    percentage = round((total_exp / monthly_income) * 100, 2) if monthly_income > 0 else 0
    savings = monthly_income - total_exp
    health_score = max(0, min(100, int(100 - percentage)))

    highest_cat = cat_expenses[0]['category'] if cat_expenses else "N/A"
    highest_amt = float(cat_expenses[0]['total']) if cat_expenses else 0

    if percentage < 60:
        message = "Your spending is under control!"
        suggestion = "Try to save more and invest wisely."
        status = "good"
    else:
        message = "Your spending is too high!"
        suggestion = "Consider reducing unnecessary expenses."
        status = "bad"

    return render_template('ai_analysis.html',
                           percentage=percentage,
                           message=message,
                           suggestion=suggestion,
                           status=status,
                           total_expenses=total_exp,
                           monthly_income=monthly_income,
                           savings=savings,
                           health_score=health_score,
                           highest_cat=highest_cat,
                           highest_amt=highest_amt)

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/investments', methods=['GET', 'POST'])
def investments():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None
    if request.method == 'POST':
        asset = request.form['asset']
        name = request.form['name']
        invested = float(request.form['invested'])
        current = float(request.form['current'])
        conn.execute(
            "INSERT INTO investments (user_id, asset, investment_name, invested_amount, current_value) VALUES (?, ?, ?, ?, ?)",
            (user_id, asset, name, invested, current))
        conn.commit()
        message = "Investment Added Successfully!"

    inv_list = conn.execute(
        "SELECT * FROM investments WHERE user_id=?",
        (user_id,)).fetchall()

    total_invested = sum(float(i['invested_amount']) for i in inv_list)
    total_current = sum(float(i['current_value']) for i in inv_list)
    profit_loss = total_current - total_invested
    roi = round((profit_loss / total_invested) * 100, 2) if total_invested > 0 else 0

    asset_data = {}
    for inv in inv_list:
        asset = inv['asset']
        if asset not in asset_data:
            asset_data[asset] = 0
        asset_data[asset] += float(inv['invested_amount'])

    asset_labels = list(asset_data.keys())
    asset_amounts = list(asset_data.values())

    conn.close()
    return render_template('investments.html',
                           investments=inv_list,
                           total_invested=total_invested,
                           total_current=total_current,
                           profit_loss=profit_loss,
                           roi=roi,
                           asset_labels=asset_labels,
                           asset_amounts=asset_amounts,
                           message=message)

@app.route('/goals', methods=['GET', 'POST'])
def goals():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None

    if request.method == 'POST':
        goal_name = request.form['goal_name']
        target = float(request.form['target'])
        saved = float(request.form['saved'])
        date = request.form['date']
        priority = request.form['priority']
        conn.execute(
            "INSERT INTO goals (user_id, goal_name, target_amount, saved_amount, target_date, priority) VALUES (?, ?, ?, ?, ?, ?)",
            (user_id, goal_name, target, saved, date, priority))
        conn.commit()
        message = "Goal Added Successfully!"

    goals_list = conn.execute(
        "SELECT * FROM goals WHERE user_id=?",
        (user_id,)).fetchall()

    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row else 0

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    available_to_save = max(0, monthly_income - total_exp)

    today = datetime.now().date()

    goals_data = []
    incomplete_goals_count = 0

    for g in goals_list:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        if saved < target:
            incomplete_goals_count += 1

    for g in goals_list:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        remaining = target - saved
        percentage = round((saved / target) * 100, 1) if target > 0 else 0
        status = "Completed" if percentage >= 100 else "In Progress"

        try:
            target_date_obj = datetime.strptime(g['target_date'], '%Y-%m-%d').date()
            days_remaining = (target_date_obj - today).days
        except (ValueError, TypeError):
            days_remaining = None

        if days_remaining and days_remaining > 0 and remaining > 0:
            months_remaining = max(1, days_remaining / 30)
            monthly_required = round(remaining / months_remaining, 2)
        elif remaining <= 0:
            monthly_required = 0
        else:
            monthly_required = remaining

        if status != "Completed" and incomplete_goals_count > 0:
            ai_suggested = round(available_to_save / incomplete_goals_count, 2)
        else:
            ai_suggested = 0

        goals_data.append({
            'id': g['id'],
            'goal_name': g['goal_name'],
            'target_amount': target,
            'saved_amount': saved,
            'remaining': remaining,
            'percentage': percentage,
            'target_date': g['target_date'],
            'status': status,
            'priority': g['priority'] if g['priority'] else 'Medium',
            'days_remaining': days_remaining,
            'monthly_required': monthly_required,
            'ai_suggested': ai_suggested,
            'just_completed': percentage >= 100
        })

    total_goals = len(goals_data)
    completed = sum(1 for g in goals_data if g['status'] == 'Completed')
    conn.close()

    return render_template('goals.html',
                           goals=goals_data,
                           message=message,
                           total_goals=total_goals,
                           completed=completed,
                           monthly_income=monthly_income,
                           available_to_save=available_to_save)

@app.route('/analytics')
def analytics():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    inv_list = conn.execute(
        "SELECT * FROM investments WHERE user_id=?",
        (user_id,)).fetchall()

    total_invested = sum(float(i['invested_amount']) for i in inv_list)
    total_current = sum(float(i['current_value']) for i in inv_list)
    profit_loss = total_current - total_invested
    roi = round((profit_loss / total_invested) * 100, 2) if total_invested > 0 else 0

    asset_data = {}
    for inv in inv_list:
        asset = inv['asset']
        if asset not in asset_data:
            asset_data[asset] = 0
        asset_data[asset] += float(inv['invested_amount'])

    asset_labels = list(asset_data.keys())
    asset_amounts = list(asset_data.values())

    inv_performance = []
    for inv in inv_list:
        pl = float(inv['current_value']) - float(inv['invested_amount'])
        ret = round((pl / float(inv['invested_amount'])) * 100, 2) if float(inv['invested_amount']) > 0 else 0
        inv_performance.append({
            'name': inv['investment_name'],
            'asset': inv['asset'],
            'pl': pl,
            'ret': ret
        })

    inv_performance.sort(key=lambda x: x['ret'], reverse=True)
    top_performers = inv_performance[:3]
    low_performers = inv_performance[-3:]

    goals_list = conn.execute(
        "SELECT * FROM goals WHERE user_id=?",
        (user_id,)).fetchall()

    goals_data = []
    for g in goals_list:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        percentage = round((saved / target) * 100, 1) if target > 0 else 0
        goals_data.append({
            'goal_name': g['goal_name'],
            'target_amount': target,
            'saved_amount': saved,
            'percentage': percentage
        })

    total_goals = len(goals_data)
    completed = sum(1 for g in goals_data if g['percentage'] >= 100)

    goal_labels = [g['goal_name'] for g in goals_data]
    goal_saved = [g['saved_amount'] for g in goals_data]
    goal_target = [g['target_amount'] for g in goals_data]

    if roi > 20:
        risk_score = "High Risk — High Return"
        risk_color = "red"
    elif roi > 10:
        risk_score = "Medium Risk — Moderate Return"
        risk_color = "orange"
    else:
        risk_score = "Low Risk — Stable Return"
        risk_color = "green"

    conn.close()

    return render_template('analytics.html',
                           total_invested=total_invested,
                           total_current=total_current,
                           profit_loss=profit_loss,
                           roi=roi,
                           asset_labels=asset_labels,
                           asset_amounts=asset_amounts,
                           top_performers=top_performers,
                           low_performers=low_performers,
                           goals_data=goals_data,
                           total_goals=total_goals,
                           completed=completed,
                           goal_labels=goal_labels,
                           goal_saved=goal_saved,
                           goal_target=goal_target,
                           risk_score=risk_score,
                           risk_color=risk_color)

@app.route('/profile', methods=['GET', 'POST'])
def profile():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None
    error = None

    if request.method == 'POST':
        current_password = request.form['current_password']
        new_password = request.form['new_password']
        confirm_password = request.form['confirm_password']

        user = conn.execute(
            "SELECT * FROM users WHERE id=?", (user_id,)).fetchone()

        if not check_password_hash(user['password'], current_password):
            error = "Current password is incorrect."
        elif new_password != confirm_password:
            error = "New passwords do not match."
        elif len(new_password) < 4:
            error = "New password must be at least 4 characters."
        else:
            conn.execute(
                "UPDATE users SET password=? WHERE id=?",
                (generate_password_hash(new_password), user_id))

    user = conn.execute(
        "SELECT * FROM users WHERE id=?", (user_id,)).fetchone()

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_inv = conn.execute(
        "SELECT SUM(invested_amount) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_goals = conn.execute(
        "SELECT COUNT(*) FROM goals WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0

    conn.close()

    return render_template('profile.html',
                           user=user,
                           message=message,
                           error=error,
                           total_exp=float(total_exp),
                           total_inv=float(total_inv),
                           total_goals=total_goals)
@app.route('/settings', methods=['GET', 'POST'])
def settings():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None

    if request.method == 'POST':
        currency = request.form['currency']
        default_income = float(request.form['default_income'] or 0)
        notification_pref = request.form.get('notification_pref', 'off')
        theme = request.form.get('theme', 'light')

        existing = conn.execute(
            "SELECT * FROM settings WHERE user_id=?", (user_id,)).fetchone()

        if existing:
            conn.execute(
                "UPDATE settings SET currency=?, default_income=?, notification_pref=?, theme=? WHERE user_id=?",
                (currency, default_income, notification_pref, theme, user_id))
        else:
            conn.execute(
                "INSERT INTO settings (user_id, currency, default_income, notification_pref, theme) VALUES (?, ?, ?, ?, ?)",
                (user_id, currency, default_income, notification_pref, theme))
        conn.commit()
        message = "Settings saved successfully!"

    user_settings = conn.execute(
        "SELECT * FROM settings WHERE user_id=?", (user_id,)).fetchone()
    user = conn.execute(
        "SELECT * FROM users WHERE id=?", (user_id,)).fetchone()
    conn.close()

    return render_template('settings.html',
                           message=message,
                           settings=user_settings,
                           user=user)
@app.route('/delete_expense/<int:expense_id>')
def delete_expense(expense_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM expenses WHERE id=? AND user_id=?",
        (expense_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('dashboard'))

@app.route('/edit_goal/<int:goal_id>', methods=['GET', 'POST'])
def edit_goal(goal_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    goal = conn.execute(
        "SELECT * FROM goals WHERE id=? AND user_id=?",
        (goal_id, user_id)).fetchone()

    if not goal:
        conn.close()
        return redirect(url_for('goals'))

    if request.method == 'POST':
        goal_name = request.form['goal_name']
        target = float(request.form['target'])
        saved = float(request.form['saved'])
        date = request.form['date']
        priority = request.form['priority']
        conn.execute(
            "UPDATE goals SET goal_name=?, target_amount=?, saved_amount=?, target_date=?, priority=? WHERE id=? AND user_id=?",
            (goal_name, target, saved, date, priority, goal_id, user_id))
        conn.commit()
        conn.close()
        return redirect(url_for('goals'))

    conn.close()
    return render_template('edit_goal.html', goal=goal)

@app.route('/delete_goal/<int:goal_id>')
def delete_goal(goal_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM goals WHERE id=? AND user_id=?",
        (goal_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('goals'))

@app.route('/edit_investment/<int:investment_id>', methods=['GET', 'POST'])
def edit_investment(investment_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    inv = conn.execute(
        "SELECT * FROM investments WHERE id=? AND user_id=?",
        (investment_id, user_id)).fetchone()

    if not inv:
        conn.close()
        return redirect(url_for('investments'))

    if request.method == 'POST':
        asset = request.form['asset']
        name = request.form['name']
        invested = float(request.form['invested'])
        current = float(request.form['current'])
        conn.execute(
            "UPDATE investments SET asset=?, investment_name=?, invested_amount=?, current_value=? WHERE id=? AND user_id=?",
            (asset, name, invested, current, investment_id, user_id))
        conn.commit()
        conn.close()
        return redirect(url_for('investments'))

    conn.close()
    return render_template('edit_investment.html', inv=inv)

@app.route('/delete_investment/<int:investment_id>')
def delete_investment(investment_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM investments WHERE id=? AND user_id=?",
        (investment_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('investments'))

@app.route('/edit_budget/<int:budget_id>', methods=['GET', 'POST'])
def edit_budget(budget_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    b = conn.execute(
        "SELECT * FROM budget WHERE id=? AND user_id=?",
        (budget_id, user_id)).fetchone()

    if not b:
        conn.close()
        return redirect(url_for('budget'))

    if request.method == 'POST':
        category = request.form['category']
        amount = float(request.form['amount'])
        month = request.form['month']
        conn.execute(
            "UPDATE budget SET category=?, amount=?, month=? WHERE id=? AND user_id=?",
            (category, amount, month, budget_id, user_id))
        conn.commit()
        conn.close()
        return redirect(url_for('budget'))

    conn.close()
    return render_template('edit_budget.html', b=b)

@app.route('/delete_budget/<int:budget_id>')
def delete_budget(budget_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM budget WHERE id=? AND user_id=?",
        (budget_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('budget'))

@app.route('/export/csv')
def export_csv():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    expenses, budgets, investments, goals = get_full_report_data(user_id)

    output = io.StringIO()
    writer = csv.writer(output)

    writer.writerow(['--- EXPENSES ---'])
    writer.writerow(['Category', 'Amount', 'Date', 'Description'])
    for e in expenses:
        writer.writerow([e['category'], e['amount'], e['date'], e['description']])

    writer.writerow([])
    writer.writerow(['--- BUDGET ---'])
    writer.writerow(['Category', 'Amount', 'Month'])
    for b in budgets:
        writer.writerow([b['category'], b['amount'], b['month']])

    writer.writerow([])
    writer.writerow(['--- INVESTMENTS ---'])
    writer.writerow(['Asset', 'Name', 'Invested', 'Current Value'])
    for i in investments:
        writer.writerow([i['asset'], i['investment_name'], i['invested_amount'], i['current_value']])

    writer.writerow([])
    writer.writerow(['--- GOALS ---'])
    writer.writerow(['Goal Name', 'Target', 'Saved', 'Target Date'])
    for g in goals:
        writer.writerow([g['goal_name'], g['target_amount'], g['saved_amount'], g['target_date']])

    mem = io.BytesIO()
    mem.write(output.getvalue().encode('utf-8'))
    mem.seek(0)
    output.close()

    return send_file(mem, mimetype='text/csv', as_attachment=True,
                     download_name='smart_finance_report.csv')

@app.route('/export/excel')
def export_excel():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    expenses, budgets, investments, goals = get_full_report_data(user_id)

    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Expenses"
    ws1.append(['Category', 'Amount', 'Date', 'Description'])
    for e in expenses:
        ws1.append([e['category'], e['amount'], e['date'], e['description']])

    ws2 = wb.create_sheet("Budget")
    ws2.append(['Category', 'Amount', 'Month'])
    for b in budgets:
        ws2.append([b['category'], b['amount'], b['month']])

    ws3 = wb.create_sheet("Investments")
    ws3.append(['Asset', 'Name', 'Invested', 'Current Value'])
    for i in investments:
        ws3.append([i['asset'], i['investment_name'], i['invested_amount'], i['current_value']])

    ws4 = wb.create_sheet("Goals")
    ws4.append(['Goal Name', 'Target', 'Saved', 'Target Date'])
    for g in goals:
        ws4.append([g['goal_name'], g['target_amount'], g['saved_amount'], g['target_date']])

    mem = io.BytesIO()
    wb.save(mem)
    mem.seek(0)

    return send_file(mem, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='smart_finance_report.xlsx')

@app.route('/export/word')
def export_word():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    expenses, budgets, investments, goals = get_full_report_data(user_id)

    doc = Document()
    doc.add_heading('Smart Finance Insights - Full Report', 0)
    doc.add_paragraph(f"Generated for: {session['user_name']}")

    doc.add_heading('Expenses', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Category', 'Amount', 'Date', 'Description'
    for e in expenses:
        row = table.add_row().cells
        row[0].text = str(e['category'])
        row[1].text = str(e['amount'])
        row[2].text = str(e['date'])
        row[3].text = str(e['description'])

    doc.add_heading('Budget', level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = 'Category', 'Amount', 'Month'
    for b in budgets:
        row = table.add_row().cells
        row[0].text = str(b['category'])
        row[1].text = str(b['amount'])
        row[2].text = str(b['month'])

    doc.add_heading('Investments', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Asset', 'Name', 'Invested', 'Current Value'
    for i in investments:
        row = table.add_row().cells
        row[0].text = str(i['asset'])
        row[1].text = str(i['investment_name'])
        row[2].text = str(i['invested_amount'])
        row[3].text = str(i['current_value'])

    doc.add_heading('Goals', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = 'Goal Name', 'Target', 'Saved', 'Target Date'
    for g in goals:
        row = table.add_row().cells
        row[0].text = str(g['goal_name'])
        row[1].text = str(g['target_amount'])
        row[2].text = str(g['saved_amount'])
        row[3].text = str(g['target_date'])

    mem = io.BytesIO()
    doc.save(mem)
    mem.seek(0)

    return send_file(mem, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                     as_attachment=True, download_name='smart_finance_report.docx')
@app.route('/spending_analysis')
def spending_analysis():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    # Category wise expenses
    cat_data = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category ORDER BY total DESC",
        (user_id,)).fetchall()

    # Total expenses
    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    # Monthly budget
    total_budget = conn.execute(
        "SELECT SUM(amount) FROM budget WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_budget = float(total_budget)
    remaining_budget = total_budget - total_exp

    # Monthly trend
    monthly_data = conn.execute(
        "SELECT strftime('%m', date) as month, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY month ORDER BY month",
        (user_id,)).fetchall()

    # Highest spending category
    highest_cat = cat_data[0]['category'] if cat_data else 'N/A'
    highest_amt = float(cat_data[0]['total']) if cat_data else 0

    conn.close()

    # Category status based on percentage
    cat_summary = []
    for cat in cat_data:
        amt = float(cat['total'])
        pct = round((amt / total_exp) * 100, 1) if total_exp > 0 else 0
        if pct >= 30:
            status = 'High'
        elif pct >= 15:
            status = 'Medium'
        else:
            status = 'Low'
        cat_summary.append({
            'category': cat['category'],
            'amount': amt,
            'percentage': pct,
            'status': status
        })

    # Chart data
    cat_labels = [c['category'] for c in cat_summary]
    cat_amounts = [c['amount'] for c in cat_summary]

    month_names = {
        '01':'Jan','02':'Feb','03':'Mar',
        '04':'Apr','05':'May','06':'Jun',
        '07':'Jul','08':'Aug','09':'Sep',
        '10':'Oct','11':'Nov','12':'Dec'
    }
    monthly_labels = [month_names.get(m['month'], m['month']) for m in monthly_data]
    monthly_amounts = [float(m['total']) for m in monthly_data]

    # Insights
    budget_utilization = round((total_exp / total_budget) * 100, 1) if total_budget > 0 else 0

    return render_template('spending_analysis.html',
                           total_exp=total_exp,
                           total_budget=total_budget,
                           remaining_budget=remaining_budget,
                           highest_cat=highest_cat,
                           highest_amt=highest_amt,
                           cat_summary=cat_summary,
                           cat_labels=cat_labels,
                           cat_amounts=cat_amounts,
                           monthly_labels=monthly_labels,
                           monthly_amounts=monthly_amounts,
                           budget_utilization=budget_utilization)
@app.route('/budget_recommendations')
def budget_recommendations():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    # Get expenses by category
    cat_expenses = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category",
        (user_id,)).fetchall()

    # Get budgets by category
    cat_budgets = conn.execute(
        "SELECT category, SUM(amount) as total FROM budget WHERE user_id=? GROUP BY category",
        (user_id,)).fetchall()

    # Income
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row and income_row['amount'] else 0

    conn.close()

    # Build comparison data
    budget_dict = {b['category']: float(b['total']) for b in cat_budgets}
    expense_dict = {e['category']: float(e['total']) for e in cat_expenses}

    all_categories = set(list(budget_dict.keys()) + list(expense_dict.keys()))

    recommendations = []
    alerts = []

    for cat in all_categories:
        spent = expense_dict.get(cat, 0)
        budget = budget_dict.get(cat, 0)
        diff = spent - budget
        pct = round((spent / budget) * 100, 1) if budget > 0 else 0

        if diff > 0:
            alerts.append({
                'category': cat,
                'spent': spent,
                'budget': budget,
                'excess': diff,
                'type': 'overspent'
            })
            recommendations.append({
                'category': cat,
                'message': f"Reduce {cat} expenses by ₹{diff:,.0f} to stay within budget.",
                'priority': 'High',
                'icon': '⚠️'
            })
        elif pct > 75:
            recommendations.append({
                'category': cat,
                'message': f"You have used {pct}% of your {cat} budget. Spend carefully.",
                'priority': 'Medium',
                'icon': '🔶'
            })

    # Savings recommendation
    total_exp = sum(expense_dict.values())
    savings = monthly_income - total_exp
    savings_ratio = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0

    if savings_ratio < 20:
        recommendations.append({
            'category': 'Savings',
            'message': f"Your savings rate is only {savings_ratio}%. Try to save at least 20% of your income.",
            'priority': 'High',
            'icon': '💰'
        })
    else:
        recommendations.append({
            'category': 'Savings',
            'message': f"Great! Your savings rate is {savings_ratio}%. Keep it up!",
            'priority': 'Low',
            'icon': '✅'
        })

    # Build chart data
    comparison_labels = list(all_categories)
    comparison_spent = [expense_dict.get(c, 0) for c in comparison_labels]
    comparison_budget = [budget_dict.get(c, 0) for c in comparison_labels]

    return render_template('budget_recommendations.html',
                           recommendations=recommendations,
                           alerts=alerts,
                           monthly_income=monthly_income,
                           total_exp=total_exp,
                           savings=savings,
                           savings_ratio=savings_ratio,
                           comparison_labels=comparison_labels,
                           comparison_spent=comparison_spent,
                           comparison_budget=comparison_budget)
@app.route('/health_score')
def health_score():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    # Income
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row and income_row['amount'] else 0

    # Expenses
    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    # Investments
    total_invested = conn.execute(
        "SELECT SUM(invested_amount) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_invested = float(total_invested)

    total_current = conn.execute(
        "SELECT SUM(current_value) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_current = float(total_current)

    # Goals
    goals = conn.execute(
        "SELECT * FROM goals WHERE user_id=?",
        (user_id,)).fetchall()
    total_goals = len(goals)
    completed_goals = sum(
        1 for g in goals
        if float(g['target_amount']) > 0 and
        float(g['saved_amount']) >= float(g['target_amount'])
    )

    conn.close()

    # Calculations
    savings = monthly_income - total_exp
    savings_ratio = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0
    expense_ratio = round((total_exp / monthly_income) * 100, 1) if monthly_income > 0 else 0
    investment_growth = round(((total_current - total_invested) / total_invested) * 100, 1) if total_invested > 0 else 0
    goal_completion = round((completed_goals / total_goals) * 100, 1) if total_goals > 0 else 0

    # Score calculation — out of 100
    score = 0

    # Savings ratio — max 30 points
    if savings_ratio >= 30:
        score += 30
    elif savings_ratio >= 20:
        score += 20
    elif savings_ratio >= 10:
        score += 10
    else:
        score += 0

    # Expense ratio — max 30 points
    if expense_ratio <= 50:
        score += 30
    elif expense_ratio <= 65:
        score += 20
    elif expense_ratio <= 80:
        score += 10
    else:
        score += 0

    # Investment growth — max 20 points
    if investment_growth >= 15:
        score += 20
    elif investment_growth >= 10:
        score += 15
    elif investment_growth >= 5:
        score += 10
    elif investment_growth > 0:
        score += 5
    else:
        score += 0

    # Goal completion — max 20 points
    if goal_completion >= 75:
        score += 20
    elif goal_completion >= 50:
        score += 15
    elif goal_completion >= 25:
        score += 10
    else:
        score += 5

    # Status
    if score >= 80:
        status = 'Excellent'
        status_color = '#4CAF50'
        status_msg = 'Your financial health is excellent! Keep it up.'
    elif score >= 60:
        status = 'Good'
        status_color = '#2196F3'
        status_msg = 'Your financial health is good. Small improvements can make it excellent.'
    elif score >= 40:
        status = 'Fair'
        status_color = '#FF9800'
        status_msg = 'Your financial health is fair. Focus on saving more and reducing expenses.'
    else:
        status = 'Poor'
        status_color = '#f44336'
        status_msg = 'Your financial health needs attention. Start by reducing unnecessary expenses.'

    # AI Recommendations
    ai_recommendations = []
    if savings_ratio < 20:
        ai_recommendations.append('Increase monthly savings to at least 20% of your income.')
    if expense_ratio > 70:
        ai_recommendations.append('Reduce your monthly expenses — currently too high.')
    if investment_growth < 10:
        ai_recommendations.append('Consider increasing your SIP by ₹2,000 per month.')
    if goal_completion < 50:
        ai_recommendations.append('Focus on completing your financial goals on time.')
    ai_recommendations.append('Maintain savings above 30% of income for financial security.')
    ai_recommendations.append('Build an emergency fund covering 6 months of expenses.')
    ai_recommendations.append('Continue maintaining a low debt ratio.')

    return render_template('health_score.html',
                           score=score,
                           status=status,
                           status_color=status_color,
                           status_msg=status_msg,
                           savings_ratio=savings_ratio,
                           expense_ratio=expense_ratio,
                           investment_growth=investment_growth,
                           goal_completion=goal_completion,
                           monthly_income=monthly_income,
                           total_exp=total_exp,
                           savings=savings,
                           total_invested=total_invested,
                           total_current=total_current,
                           ai_recommendations=ai_recommendations)
@app.route('/notifications')
def notifications():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    # Auto generate notifications
    # Clear old auto notifications first
    conn.execute(
        "DELETE FROM notifications WHERE user_id=? AND type != 'Bill Reminder'",
        (user_id,))
    conn.commit()

    # Income
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row and income_row['amount'] else 0

    # Expenses by category
    cat_expenses = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category",
        (user_id,)).fetchall()

    # Budgets by category
    cat_budgets = conn.execute(
        "SELECT category, SUM(amount) as total FROM budget WHERE user_id=? GROUP BY category",
        (user_id,)).fetchall()

    budget_dict = {b['category']: float(b['total']) for b in cat_budgets}
    expense_dict = {e['category']: float(e['total']) for e in cat_expenses}

    from datetime import date
    today = date.today().strftime('%Y-%m-%d')

    # Budget alerts
    for cat, spent in expense_dict.items():
        budget = budget_dict.get(cat, 0)
        if budget > 0 and spent > budget:
            excess = spent - budget
            conn.execute(
                "INSERT INTO notifications (user_id, type, message, priority, status, created_date) VALUES (?, ?, ?, ?, ?, ?)",
                (user_id, '⚠️ Budget Alert',
                 f"{cat} expenses exceeded budget by ₹{excess:,.0f}",
                 'High', 'Active', today))
        elif budget > 0 and (spent / budget) >= 0.8:
            conn.execute(
                "INSERT INTO notifications (user_id, type, message, priority, status, created_date) VALUES (?, ?, ?, ?, ?, ?)",
                (user_id, '🔶 Budget Warning',
                 f"{cat} budget is 80% used. Spend carefully.",
                 'Medium', 'Active', today))

    # Total expenses
    total_exp = sum(expense_dict.values())
    savings = monthly_income - total_exp
    savings_ratio = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0

    # Low savings alert
    if savings_ratio < 20 and monthly_income > 0:
        conn.execute(
            "INSERT INTO notifications (user_id, type, message, priority, status, created_date) VALUES (?, ?, ?, ?, ?, ?)",
            (user_id, '💰 Savings Alert',
             f"Your savings rate is only {savings_ratio}%. Try to save at least 20% of your income.",
             'High', 'Active', today))

    # Goal reminders
    goals = conn.execute(
        "SELECT * FROM goals WHERE user_id=?",
        (user_id,)).fetchall()

    for g in goals:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        remaining = target - saved
        pct = round((saved / target) * 100, 1) if target > 0 else 0

        if pct < 100:
            conn.execute(
                "INSERT INTO notifications (user_id, type, message, priority, status, created_date) VALUES (?, ?, ?, ?, ?, ?)",
                (user_id, '🎯 Goal Reminder',
                 f"Save ₹{remaining:,.0f} more to reach your {g['goal_name']} goal ({pct}% complete).",
                 'Medium', 'Active', today))

    # Investment alerts
    investments = conn.execute(
        "SELECT * FROM investments WHERE user_id=?",
        (user_id,)).fetchall()

    for inv in investments:
        pl = float(inv['current_value']) - float(inv['invested_amount'])
        roi = round((pl / float(inv['invested_amount'])) * 100, 2) if float(inv['invested_amount']) > 0 else 0
        if roi > 10:
            conn.execute(
                "INSERT INTO notifications (user_id, type, message, priority, status, created_date) VALUES (?, ?, ?, ?, ?, ?)",
                (user_id, '📈 Investment Alert',
                 f"{inv['investment_name']} portfolio increased by {roi}%. Great performance!",
                 'Low', 'Completed', today))
        elif roi < 0:
            conn.execute(
                "INSERT INTO notifications (user_id, type, message, priority, status, created_date) VALUES (?, ?, ?, ?, ?, ?)",
                (user_id, '📉 Investment Alert',
                 f"{inv['investment_name']} is at a loss of {abs(roi)}%. Consider reviewing.",
                 'High', 'Active', today))

    conn.commit()

    # Fetch all notifications
    all_notifs = conn.execute(
        "SELECT * FROM notifications WHERE user_id=? ORDER BY id DESC",
        (user_id,)).fetchall()

    # Count by priority
    high = sum(1 for n in all_notifs if n['priority'] == 'High')
    medium = sum(1 for n in all_notifs if n['priority'] == 'Medium')
    low = sum(1 for n in all_notifs if n['priority'] == 'Low')

    conn.close()

    return render_template('notifications.html',
                           notifications=all_notifs,
                           total=len(all_notifs),
                           high=high,
                           medium=medium,
                           low=low)
@app.route('/ai_insights')
def ai_insights():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    # Income
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row and income_row['amount'] else 0

    # Expenses
    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    # Category expenses
    cat_expenses = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category ORDER BY total DESC",
        (user_id,)).fetchall()

    # Monthly trend
    monthly_data = conn.execute(
        "SELECT strftime('%m', date) as month, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY month ORDER BY month",
        (user_id,)).fetchall()

    # Investments
    total_invested = conn.execute(
        "SELECT SUM(invested_amount) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_invested = float(total_invested)

    total_current = conn.execute(
        "SELECT SUM(current_value) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_current = float(total_current)

    # Goals
    goals = conn.execute(
        "SELECT * FROM goals WHERE user_id=?",
        (user_id,)).fetchall()

    conn.close()

    # Calculations
    savings = monthly_income - total_exp
    savings_ratio = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0
    expense_ratio = round((total_exp / monthly_income) * 100, 1) if monthly_income > 0 else 0
    investment_growth = round(((total_current - total_invested) / total_invested) * 100, 1) if total_invested > 0 else 0

    # Monthly trend data
    month_names = {
        '01':'Jan','02':'Feb','03':'Mar',
        '04':'Apr','05':'May','06':'Jun',
        '07':'Jul','08':'Aug','09':'Sep',
        '10':'Oct','11':'Nov','12':'Dec'
    }
    monthly_labels = [month_names.get(m['month'], m['month']) for m in monthly_data]
    monthly_amounts = [float(m['total']) for m in monthly_data]

    # Category data
    cat_labels = [c['category'] for c in cat_expenses]
    cat_amounts = [float(c['total']) for c in cat_expenses]

    # AI Insights generation
    spending_insights = []
    saving_insights = []
    investment_insights = []
    goal_insights = []

    # Spending insights
    if cat_expenses:
        top_cat = cat_expenses[0]['category']
        top_amt = float(cat_expenses[0]['total'])
        top_pct = round((top_amt / total_exp) * 100, 1) if total_exp > 0 else 0
        spending_insights.append({
            'icon': '📊',
            'title': 'Top Spending Category',
            'message': f"You spend the most on {top_cat} — ₹{top_amt:,.0f} ({top_pct}% of total expenses).",
            'type': 'warning' if top_pct > 30 else 'info'
        })

    if expense_ratio > 70:
        spending_insights.append({
            'icon': '⚠️',
            'title': 'High Expense Ratio',
            'message': f"You are spending {expense_ratio}% of your income. Try to keep it below 70%.",
            'type': 'danger'
        })
    else:
        spending_insights.append({
            'icon': '✅',
            'title': 'Healthy Expense Ratio',
            'message': f"Your expense ratio is {expense_ratio}%. You are managing expenses well.",
            'type': 'success'
        })

    if len(monthly_data) >= 2:
        last = float(monthly_data[-1]['total'])
        prev = float(monthly_data[-2]['total'])
        diff = last - prev
        if diff > 0:
            spending_insights.append({
                'icon': '📈',
                'title': 'Spending Increased',
                'message': f"Your spending increased by ₹{diff:,.0f} compared to last month. Review your expenses.",
                'type': 'warning'
            })
        else:
            spending_insights.append({
                'icon': '📉',
                'title': 'Spending Decreased',
                'message': f"Great! Your spending decreased by ₹{abs(diff):,.0f} compared to last month.",
                'type': 'success'
            })

    # Saving insights
    if savings_ratio >= 30:
        saving_insights.append({
            'icon': '🌟',
            'title': 'Excellent Savings Rate',
            'message': f"You are saving {savings_ratio}% of your income. Outstanding financial discipline!",
            'type': 'success'
        })
    elif savings_ratio >= 20:
        saving_insights.append({
            'icon': '✅',
            'title': 'Good Savings Rate',
            'message': f"You are saving {savings_ratio}% of your income. Try to reach 30% for better security.",
            'type': 'info'
        })
    else:
        saving_insights.append({
            'icon': '⚠️',
            'title': 'Low Savings Rate',
            'message': f"You are only saving {savings_ratio}% of income. Increase savings to at least 20%.",
            'type': 'danger'
        })

    saving_insights.append({
        'icon': '💡',
        'title': 'Savings Tip',
        'message': "Automate your savings by setting up a monthly transfer on salary day. Pay yourself first!",
        'type': 'info'
    })

    saving_insights.append({
        'icon': '🛡️',
        'title': 'Emergency Fund',
        'message': f"Build an emergency fund of ₹{(total_exp * 6):,.0f} — enough to cover 6 months of expenses.",
        'type': 'info'
    })

    # Investment insights
    if total_invested > 0:
        if investment_growth > 15:
            investment_insights.append({
                'icon': '🚀',
                'title': 'Strong Investment Performance',
                'message': f"Your portfolio grew by {investment_growth}%. Excellent returns!",
                'type': 'success'
            })
        elif investment_growth > 0:
            investment_insights.append({
                'icon': '📈',
                'title': 'Positive Investment Growth',
                'message': f"Your portfolio grew by {investment_growth}%. Consider increasing SIP by ₹2,000.",
                'type': 'info'
            })
        else:
            investment_insights.append({
                'icon': '📉',
                'title': 'Investment Loss',
                'message': f"Your portfolio is down by {abs(investment_growth)}%. Review your investment strategy.",
                'type': 'danger'
            })

        investment_insights.append({
            'icon': '💼',
            'title': 'Diversification Tip',
            'message': "Diversify across Stocks, Mutual Funds, Gold, and Fixed Deposits to reduce risk.",
            'type': 'info'
        })
    else:
        investment_insights.append({
            'icon': '💡',
            'title': 'Start Investing',
            'message': "You have no investments yet. Start a SIP of ₹2,000/month to grow your wealth.",
            'type': 'warning'
        })

    # Goal insights
    for g in goals:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        pct = round((saved / target) * 100, 1) if target > 0 else 0
        remaining = target - saved
        if pct >= 100:
            goal_insights.append({
                'icon': '🎉',
                'title': f"{g['goal_name']} — Completed!",
                'message': "Congratulations! You have achieved this financial goal.",
                'type': 'success'
            })
        elif pct >= 75:
            goal_insights.append({
                'icon': '🎯',
                'title': f"{g['goal_name']} — Almost There!",
                'message': f"You need ₹{remaining:,.0f} more to complete this goal. Keep going!",
                'type': 'info'
            })
        else:
            goal_insights.append({
                'icon': '📌',
                'title': f"{g['goal_name']} — In Progress",
                'message': f"{pct}% complete. Save ₹{remaining:,.0f} more to reach your target.",
                'type': 'warning'
            })

    if not goals:
        goal_insights.append({
            'icon': '💡',
            'title': 'Set Financial Goals',
            'message': "You have no goals yet. Set goals like Emergency Fund, Vacation, or Home Purchase.",
            'type': 'info'
        })

    return render_template('ai_insights.html',
                           monthly_income=monthly_income,
                           total_exp=total_exp,
                           savings=savings,
                           savings_ratio=savings_ratio,
                           expense_ratio=expense_ratio,
                           investment_growth=investment_growth,
                           spending_insights=spending_insights,
                           saving_insights=saving_insights,
                           investment_insights=investment_insights,
                           goal_insights=goal_insights,
                           monthly_labels=monthly_labels,
                           monthly_amounts=monthly_amounts,
                           cat_labels=cat_labels,
                           cat_amounts=cat_amounts)
@app.route('/intelligence_dashboard')
def intelligence_dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    # Income
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row and income_row['amount'] else 0

    # Expenses
    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    # Category expenses
    cat_expenses = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category ORDER BY total DESC",
        (user_id,)).fetchall()

    # Budget
    total_budget = conn.execute(
        "SELECT SUM(amount) FROM budget WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_budget = float(total_budget)

    cat_budgets = conn.execute(
        "SELECT category, SUM(amount) as total FROM budget WHERE user_id=? GROUP BY category",
        (user_id,)).fetchall()

    # Investments
    total_invested = conn.execute(
        "SELECT SUM(invested_amount) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_invested = float(total_invested)

    total_current = conn.execute(
        "SELECT SUM(current_value) FROM investments WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_current = float(total_current)

    # Goals
    goals = conn.execute(
        "SELECT * FROM goals WHERE user_id=?",
        (user_id,)).fetchall()

    # Notifications
    notifications = conn.execute(
        "SELECT * FROM notifications WHERE user_id=? AND priority='High' ORDER BY id DESC LIMIT 5",
        (user_id,)).fetchall()

    # Monthly trend
    monthly_data = conn.execute(
        "SELECT strftime('%m', date) as month, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY month ORDER BY month",
        (user_id,)).fetchall()

    conn.close()

    # Calculations
    savings = monthly_income - total_exp
    savings_ratio = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0
    expense_ratio = round((total_exp / monthly_income) * 100, 1) if monthly_income > 0 else 0
    investment_growth = round(((total_current - total_invested) / total_invested) * 100, 1) if total_invested > 0 else 0
    budget_utilization = round((total_exp / total_budget) * 100, 1) if total_budget > 0 else 0
    remaining_budget = total_budget - total_exp

    # Health Score
    score = 0
    savings_score = min(30, int((savings_ratio / 30) * 30))
    expense_score = max(0, 30 - int((expense_ratio / 100) * 30))
    investment_score = min(20, int((investment_growth / 15) * 20)) if investment_growth > 0 else 0
    goal_pcts = [float(g['saved_amount']) / float(g['target_amount']) * 100 for g in goals if float(g['target_amount']) > 0]
    avg_goal = sum(goal_pcts) / len(goal_pcts) if goal_pcts else 0
    goal_score = min(20, int((avg_goal / 100) * 20))
    score = savings_score + expense_score + investment_score + goal_score

    if score >= 80:
        status = 'Excellent'
        status_color = '#4CAF50'
    elif score >= 60:
        status = 'Good'
        status_color = '#2196F3'
    elif score >= 40:
        status = 'Fair'
        status_color = '#FF9800'
    else:
        status = 'Poor'
        status_color = '#f44336'

    # Budget recommendations
    budget_dict = {b['category']: float(b['total']) for b in cat_budgets}
    expense_dict = {e['category']: float(e['total']) for e in cat_expenses}
    recommendations = []
    for cat in expense_dict:
        spent = expense_dict[cat]
        budget = budget_dict.get(cat, 0)
        if budget > 0 and spent > budget:
            recommendations.append(f"Reduce {cat} by ₹{(spent-budget):,.0f}")
    if savings_ratio < 20:
        recommendations.append(f"Increase monthly savings by ₹{(monthly_income * 0.2 - savings):,.0f}")

    # Goals data
    goals_data = []
    for g in goals:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        pct = round((saved / target) * 100, 1) if target > 0 else 0
        goals_data.append({
            'goal_name': g['goal_name'],
            'percentage': pct,
            'saved': saved,
            'target': target
        })

    # Chart data
    cat_labels = [c['category'] for c in cat_expenses]
    cat_amounts = [float(c['total']) for c in cat_expenses]

    month_names = {
        '01':'Jan','02':'Feb','03':'Mar',
        '04':'Apr','05':'May','06':'Jun',
        '07':'Jul','08':'Aug','09':'Sep',
        '10':'Oct','11':'Nov','12':'Dec'
    }
    monthly_labels = [month_names.get(m['month'], m['month']) for m in monthly_data]
    monthly_amounts = [float(m['total']) for m in monthly_data]

    # AI Insights summary
    ai_insights_list = []
    if savings_ratio >= 20:
        ai_insights_list.append(f"✅ Savings rate is healthy at {savings_ratio}%.")
    else:
        ai_insights_list.append(f"⚠️ Savings rate is low at {savings_ratio}%. Increase savings.")

    if investment_growth > 0:
        ai_insights_list.append(f"📈 Investment portfolio gained {investment_growth}% this month.")
    else:
        ai_insights_list.append("💡 Consider reviewing your investment portfolio.")

    if total_budget > 0:
        months_covered = round(savings * 6 / monthly_income, 1) if monthly_income > 0 else 0
        ai_insights_list.append(f"🛡️ Emergency fund covers approximately {months_covered} months of expenses.")

    return render_template('intelligence_dashboard.html',
                           monthly_income=monthly_income,
                           total_exp=total_exp,
                           savings=savings,
                           savings_ratio=savings_ratio,
                           expense_ratio=expense_ratio,
                           investment_growth=investment_growth,
                           total_invested=total_invested,
                           total_current=total_current,
                           budget_utilization=budget_utilization,
                           remaining_budget=remaining_budget,
                           score=score,
                           status=status,
                           status_color=status_color,
                           recommendations=recommendations,
                           goals_data=goals_data,
                           notifications=notifications,
                           cat_labels=cat_labels,
                           cat_amounts=cat_amounts,
                           monthly_labels=monthly_labels,
                           monthly_amounts=monthly_amounts,
                           ai_insights_list=ai_insights_list)

@app.route('/wallets', methods=['GET', 'POST'])
def wallets():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None

    if request.method == 'POST':
        wallet_name = request.form['wallet_name']
        wallet_type = request.form['wallet_type']
        balance = float(request.form['balance'])
        conn.execute(
            "INSERT INTO wallets (user_id, wallet_name, wallet_type, balance) VALUES (?, ?, ?, ?)",
            (user_id, wallet_name, wallet_type, balance))
        conn.commit()
        message = "Wallet Added Successfully!"

    wallet_list = conn.execute(
        "SELECT * FROM wallets WHERE user_id=?", (user_id,)).fetchall()
    total_balance = sum(float(w['balance']) for w in wallet_list)
    conn.close()

    return render_template('wallets.html',
                           wallets=wallet_list,
                           total_balance=total_balance,
                           message=message)

@app.route('/edit_wallet/<int:wallet_id>', methods=['GET', 'POST'])
def edit_wallet(wallet_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    wallet = conn.execute(
        "SELECT * FROM wallets WHERE id=? AND user_id=?",
        (wallet_id, user_id)).fetchone()

    if not wallet:
        conn.close()
        return redirect(url_for('wallets'))

    if request.method == 'POST':
        wallet_name = request.form['wallet_name']
        wallet_type = request.form['wallet_type']
        balance = float(request.form['balance'])
        conn.execute(
            "UPDATE wallets SET wallet_name=?, wallet_type=?, balance=? WHERE id=? AND user_id=?",
            (wallet_name, wallet_type, balance, wallet_id, user_id))
        conn.commit()
        conn.close()
        return redirect(url_for('wallets'))

    conn.close()
    return render_template('edit_wallet.html', wallet=wallet)

@app.route('/delete_wallet/<int:wallet_id>')
def delete_wallet(wallet_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM wallets WHERE id=? AND user_id=?",
        (wallet_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('wallets'))
@app.route('/bills', methods=['GET', 'POST'])
def bills():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None

    if request.method == 'POST':
        bill_name = request.form['bill_name']
        amount = float(request.form['amount'])
        due_date = request.form['due_date']
        recurring = request.form['recurring']
        conn.execute(
            "INSERT INTO bills (user_id, bill_name, amount, due_date, recurring, paid) VALUES (?, ?, ?, ?, ?, 'No')",
            (user_id, bill_name, amount, due_date, recurring))
        conn.commit()
        message = "Bill Added Successfully!"

    bills_list = conn.execute(
        "SELECT * FROM bills WHERE user_id=? ORDER BY due_date ASC",
        (user_id,)).fetchall()
    conn.close()

    today = datetime.now().date()
    bills_data = []
    for b in bills_list:
        try:
            due = datetime.strptime(b['due_date'], '%Y-%m-%d').date()
            days_left = (due - today).days
        except (ValueError, TypeError):
            days_left = None

        if b['paid'] == 'Yes':
            urgency = 'paid'
        elif days_left is not None and days_left < 0:
            urgency = 'overdue'
        elif days_left is not None and days_left <= 3:
            urgency = 'urgent'
        elif days_left is not None and days_left <= 7:
            urgency = 'soon'
        else:
            urgency = 'upcoming'

        bills_data.append({
            'id': b['id'],
            'bill_name': b['bill_name'],
            'amount': float(b['amount']),
            'due_date': b['due_date'],
            'due_day': due.day if days_left is not None else None,
            'recurring': b['recurring'],
            'paid': b['paid'],
            'days_left': days_left,
            'urgency': urgency
        })

    # Build a simple calendar grid for the current month
    year = today.year
    month = today.month
    cal = calendar.Calendar(firstweekday=6)  # Sunday start
    month_days = cal.monthdayscalendar(year, month)
    month_name = calendar.month_name[month]

    bills_by_day = {}
    for b in bills_data:
        if b['due_day']:
            try:
                due_obj = datetime.strptime(b['due_date'], '%Y-%m-%d').date()
                if due_obj.year == year and due_obj.month == month:
                    bills_by_day.setdefault(due_obj.day, []).append(b['bill_name'])
            except (ValueError, TypeError):
                pass

    reminders = [b for b in bills_data if b['urgency'] in ('overdue', 'urgent', 'soon')]

    return render_template('bills.html',
                           bills=bills_data,
                           reminders=reminders,
                           month_days=month_days,
                           bills_by_day=bills_by_day,
                           month_name=month_name,
                           year=year,
                           today_day=today.day)

@app.route('/mark_paid/<int:bill_id>')
def mark_paid(bill_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "UPDATE bills SET paid='Yes' WHERE id=? AND user_id=?",
        (bill_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('bills'))

@app.route('/delete_bill/<int:bill_id>')
def delete_bill(bill_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM bills WHERE id=? AND user_id=?",
        (bill_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('bills'))
@app.route('/subscriptions', methods=['GET', 'POST'])
def subscriptions():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None

    if request.method == 'POST':
        sub_name = request.form['subscription_name']
        amount = float(request.form['amount'])
        billing_cycle = request.form['billing_cycle']
        next_date = request.form['next_billing_date']
        category = request.form['category']
        conn.execute(
            "INSERT INTO subscriptions (user_id, subscription_name, amount, billing_cycle, next_billing_date, category) VALUES (?, ?, ?, ?, ?, ?)",
            (user_id, sub_name, amount, billing_cycle, next_date, category))
        conn.commit()
        message = "Subscription Added Successfully!"

    sub_list = conn.execute(
        "SELECT * FROM subscriptions WHERE user_id=? ORDER BY next_billing_date ASC",
        (user_id,)).fetchall()
    conn.close()

    today = datetime.now().date()
    subs_data = []
    total_monthly_cost = 0

    for s in sub_list:
        amount = float(s['amount'])
        cycle = s['billing_cycle']
        monthly_equivalent = amount if cycle == 'Monthly' else round(amount / 12, 2)
        total_monthly_cost += monthly_equivalent

        try:
            next_date_obj = datetime.strptime(s['next_billing_date'], '%Y-%m-%d').date()
            days_until = (next_date_obj - today).days
        except (ValueError, TypeError):
            days_until = None

        subs_data.append({
            'id': s['id'],
            'subscription_name': s['subscription_name'],
            'amount': amount,
            'billing_cycle': cycle,
            'next_billing_date': s['next_billing_date'],
            'category': s['category'],
            'monthly_equivalent': monthly_equivalent,
            'days_until': days_until
        })

    total_yearly_cost = round(total_monthly_cost * 12, 2)

    return render_template('subscriptions.html',
                           subscriptions=subs_data,
                           total_monthly_cost=round(total_monthly_cost, 2),
                           total_yearly_cost=total_yearly_cost,
                           message=message)

@app.route('/edit_subscription/<int:sub_id>', methods=['GET', 'POST'])
def edit_subscription(sub_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    sub = conn.execute(
        "SELECT * FROM subscriptions WHERE id=? AND user_id=?",
        (sub_id, user_id)).fetchone()

    if not sub:
        conn.close()
        return redirect(url_for('subscriptions'))

    if request.method == 'POST':
        sub_name = request.form['subscription_name']
        amount = float(request.form['amount'])
        billing_cycle = request.form['billing_cycle']
        next_date = request.form['next_billing_date']
        category = request.form['category']
        conn.execute(
            "UPDATE subscriptions SET subscription_name=?, amount=?, billing_cycle=?, next_billing_date=?, category=? WHERE id=? AND user_id=?",
            (sub_name, amount, billing_cycle, next_date, category, sub_id, user_id))
        conn.commit()
        conn.close()
        return redirect(url_for('subscriptions'))

    conn.close()
    return render_template('edit_subscription.html', sub=sub)

@app.route('/delete_subscription/<int:sub_id>')
def delete_subscription(sub_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute(
        "DELETE FROM subscriptions WHERE id=? AND user_id=?",
        (sub_id, user_id))
    conn.commit()
    conn.close()
    return redirect(url_for('subscriptions'))
@app.route('/challenges', methods=['GET', 'POST'])
def challenges():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    message = None
    today = datetime.now().date()
    today_str = today.strftime('%Y-%m-%d')

    if request.method == 'POST':
        challenge_name = request.form['challenge_name']
        challenge_type = request.form['challenge_type']
        target_days = int(request.form['target_days'])
        conn.execute(
            "INSERT INTO challenges (user_id, challenge_name, challenge_type, target_days, start_date, status) VALUES (?, ?, ?, ?, ?, 'Active')",
            (user_id, challenge_name, challenge_type, target_days, today_str))
        conn.commit()
        message = "Challenge Created Successfully!"

    challenge_list = conn.execute(
        "SELECT * FROM challenges WHERE user_id=?", (user_id,)).fetchall()

    challenges_data = []
    for c in challenge_list:
        log_count = conn.execute(
            "SELECT COUNT(*) FROM challenge_logs WHERE challenge_id=?",
            (c['id'],)).fetchone()[0]
        progress = round((log_count / c['target_days']) * 100, 1) if c['target_days'] > 0 else 0
        already_logged_today = conn.execute(
            "SELECT * FROM challenge_logs WHERE challenge_id=? AND log_date=?",
            (c['id'], today_str)).fetchone() is not None
        status = 'Completed' if log_count >= c['target_days'] else c['status']

        challenges_data.append({
            'id': c['id'],
            'challenge_name': c['challenge_name'],
            'challenge_type': c['challenge_type'],
            'target_days': c['target_days'],
            'log_count': log_count,
            'progress': min(progress, 100),
            'already_logged_today': already_logged_today,
            'status': status
        })

    # --- No-Spend Streak ---
    expense_dates = conn.execute(
        "SELECT DISTINCT date FROM expenses WHERE user_id=?", (user_id,)).fetchall()
    spend_dates = set(row['date'] for row in expense_dates)

    streak = 0
    check_day = today
    for _ in range(365):
        check_str = check_day.strftime('%Y-%m-%d')
        if check_str not in spend_dates:
            streak += 1
            check_day -= timedelta(days=1)
        else:
            break

    # --- Milestones (auto-detected) ---
    milestones = []

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?", (user_id,)).fetchone()[0] or 0
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1", (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row else 0
    savings = monthly_income - float(total_exp)

    if savings >= 100:
        milestones.append("💰 First ₹100 Saved")
    if savings >= 10000:
        milestones.append("🏦 Saved Over ₹10,000")

    completed_goals = conn.execute(
        "SELECT COUNT(*) FROM goals WHERE user_id=? AND saved_amount >= target_amount",
        (user_id,)).fetchone()[0]
    if completed_goals >= 1:
        milestones.append("🎯 Completed a Financial Goal")
    if completed_goals >= 3:
        milestones.append("🏆 Completed 3 Financial Goals")

    budgets = conn.execute("SELECT * FROM budget WHERE user_id=?", (user_id,)).fetchall()
    on_track_count = 0
    for b in budgets:
        spent = conn.execute(
            "SELECT SUM(amount) FROM expenses WHERE user_id=? AND category=?",
            (user_id, b['category'])).fetchone()[0] or 0
        if float(spent) <= float(b['amount']):
            on_track_count += 1
    if on_track_count >= 3:
        milestones.append("📊 3 Budget Categories On Track")

    if streak >= 7:
        milestones.append("🔥 7-Day No-Spend Streak")
    if streak >= 30:
        milestones.append("🔥🔥 30-Day No-Spend Streak")

    completed_challenges = sum(1 for c in challenges_data if c['status'] == 'Completed')
    if completed_challenges >= 1:
        milestones.append("✅ Completed a Savings Challenge")

    conn.close()

    return render_template('challenges.html',
                           challenges=challenges_data,
                           streak=streak,
                           milestones=milestones,
                           message=message)

@app.route('/log_challenge/<int:challenge_id>')
def log_challenge(challenge_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    conn = get_db()
    today_str = datetime.now().strftime('%Y-%m-%d')
    already = conn.execute(
        "SELECT * FROM challenge_logs WHERE challenge_id=? AND log_date=?",
        (challenge_id, today_str)).fetchone()
    if not already:
        conn.execute(
            "INSERT INTO challenge_logs (challenge_id, log_date) VALUES (?, ?)",
            (challenge_id, today_str))
        conn.commit()
    conn.close()
    return redirect(url_for('challenges'))

@app.route('/delete_challenge/<int:challenge_id>')
def delete_challenge(challenge_id):
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()
    conn.execute("DELETE FROM challenges WHERE id=? AND user_id=?", (challenge_id, user_id))
    conn.execute("DELETE FROM challenge_logs WHERE challenge_id=?", (challenge_id,))
    conn.commit()
    conn.close()
    return redirect(url_for('challenges'))
@app.route('/spending_personality')
def spending_personality():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row else 0

    savings = monthly_income - total_exp
    savings_rate = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0

    # Discretionary vs essential spending
    discretionary_categories = ['Shopping', 'Entertainment', 'Others']
    discretionary_total = conn.execute(
        f"SELECT SUM(amount) FROM expenses WHERE user_id=? AND category IN ({','.join('?'*len(discretionary_categories))})",
        (user_id, *discretionary_categories)).fetchone()[0] or 0
    discretionary_total = float(discretionary_total)
    discretionary_pct = round((discretionary_total / total_exp) * 100, 1) if total_exp > 0 else 0

    # Budget adherence
    budgets = conn.execute("SELECT * FROM budget WHERE user_id=?", (user_id,)).fetchall()
    over_budget_count = 0
    for b in budgets:
        spent = conn.execute(
            "SELECT SUM(amount) FROM expenses WHERE user_id=? AND category=?",
            (user_id, b['category'])).fetchone()[0] or 0
        if float(spent) > float(b['amount']):
            over_budget_count += 1
    total_budgets = len(budgets)
    over_budget_pct = round((over_budget_count / total_budgets) * 100, 1) if total_budgets > 0 else 0

    conn.close()

    # --- Classification logic ---
    score = 0
    reasons = []

    if savings_rate >= 30:
        score += 2
        reasons.append(f"You save {savings_rate}% of your income — well above the recommended 20%.")
    elif savings_rate >= 10:
        score += 1
        reasons.append(f"You save {savings_rate}% of your income — a reasonable, moderate rate.")
    else:
        reasons.append(f"You save only {savings_rate}% of your income, or are in deficit.")

    if discretionary_pct <= 25:
        score += 2
        reasons.append(f"Only {discretionary_pct}% of spending goes to discretionary categories (Shopping/Entertainment/Others).")
    elif discretionary_pct <= 45:
        score += 1
        reasons.append(f"{discretionary_pct}% of spending is discretionary — a moderate share.")
    else:
        reasons.append(f"{discretionary_pct}% of spending is discretionary — a high share of non-essential spend.")

    if over_budget_pct <= 20:
        score += 2
        reasons.append(f"You stayed within budget in {100 - over_budget_pct:.0f}% of your categories.")
    elif over_budget_pct <= 50:
        score += 1
        reasons.append(f"You went over budget in {over_budget_pct}% of your categories.")
    else:
        reasons.append(f"You went over budget in {over_budget_pct}% of your categories — frequent overspending.")

    if score >= 5:
        personality = "Saver"
        emoji = "🐢"
        description = "You're disciplined with money — prioritizing savings and staying within budget."
        color = "green"
        tips = [
            "Consider putting extra savings into investments for better long-term returns.",
            "Keep an emergency fund of 3-6 months of expenses if you haven't already.",
        ]
    elif score >= 3:
        personality = "Balanced"
        emoji = "⚖️"
        description = "You balance spending and saving reasonably well, with some room to tighten up."
        color = "orange"
        tips = [
            "Try to identify one discretionary category to trim by 10% this month.",
            "Set a specific savings goal to give your extra income a clear purpose.",
        ]
    else:
        personality = "Frequent Spender"
        emoji = "🛍️"
        description = "Spending currently outpaces saving — there's an opportunity to build better habits."
        color = "red"
        tips = [
            "Start with the 50/30/20 rule: 50% needs, 30% wants, 20% savings.",
            "Try a short no-spend challenge to reset spending habits.",
            "Review your largest discretionary category and set a hard budget cap for it.",
        ]

    return render_template('spending_personality.html',
                           personality=personality,
                           emoji=emoji,
                           description=description,
                           color=color,
                           reasons=reasons,
                           tips=tips,
                           savings_rate=savings_rate,
                           discretionary_pct=discretionary_pct,
                           over_budget_pct=over_budget_pct)
@app.route('/simulator')
def simulator():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row else 0

    conn.close()

    current_savings = monthly_income - total_exp

    return render_template('simulator.html',
                           monthly_income=monthly_income,
                           total_expenses=total_exp,
                           current_savings=current_savings)
# --- Smart Notifications ---
    notifications = []

    today = datetime.now().date()
    week_start = today - timedelta(days=today.weekday())  # Monday of this week
    week_start_str = week_start.strftime('%Y-%m-%d')

    this_week_spend = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=? AND date >= ?",
        (user_id, week_start_str)).fetchone()[0] or 0
    this_week_spend = float(this_week_spend)

    # Average weekly spend over the past 4 completed weeks (excluding this week)
    four_weeks_ago = (week_start - timedelta(weeks=4)).strftime('%Y-%m-%d')
    past_total = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=? AND date >= ? AND date < ?",
        (user_id, four_weeks_ago, week_start_str)).fetchone()[0] or 0
    past_avg_weekly = float(past_total) / 4

    if past_avg_weekly > 0 and this_week_spend > past_avg_weekly * 1.3:
        pct_higher = round(((this_week_spend - past_avg_weekly) / past_avg_weekly) * 100)
        notifications.append({
            'type': 'warning',
            'icon': '⚠️',
            'text': f"Your spending is {pct_higher}% higher than usual this week."
        })
    elif past_avg_weekly > 0 and this_week_spend < past_avg_weekly * 0.7:
        notifications.append({
            'type': 'good',
            'icon': '👍',
            'text': "Great job — you're spending less than usual this week!"
        })

    # Category spike check
    cat_this_week = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? AND date >= ? GROUP BY category",
        (user_id, week_start_str)).fetchall()
    for row in cat_this_week:
        cat_past = conn.execute(
            "SELECT SUM(amount) FROM expenses WHERE user_id=? AND category=? AND date >= ? AND date < ?",
            (user_id, row['category'], four_weeks_ago, week_start_str)).fetchone()[0] or 0
        cat_past_avg = float(cat_past) / 4
        if cat_past_avg > 0 and float(row['total']) > cat_past_avg * 1.5:
            notifications.append({
                'type': 'warning',
                'icon': '📈',
                'text': f"{row['category']} spending is unusually high this week."
            })

    # Overdue bills check (only if bills table has data)
    try:
        overdue_bills = conn.execute(
            "SELECT COUNT(*) FROM bills WHERE user_id=? AND paid='No' AND due_date < ?",
            (user_id, today.strftime('%Y-%m-%d'))).fetchone()[0]
        if overdue_bills > 0:
            notifications.append({
                'type': 'warning',
                'icon': '🔔',
                'text': f"You have {overdue_bills} overdue bill(s)."
            })
    except sqlite3.OperationalError:
        pass
@app.route('/bill_splitter')
def bill_splitter():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('bill_splitter.html')
@app.route('/reports')
def reports():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    return render_template('reports.html')
@app.route('/reports/expense')
def report_expense():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    current_month = datetime.now().strftime('%Y-%m')

    expenses = conn.execute(
        "SELECT * FROM expenses WHERE user_id=? AND date LIKE ? ORDER BY date DESC",
        (user_id, current_month + '%')).fetchall()

    cat_totals = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? AND date LIKE ? GROUP BY category ORDER BY total DESC",
        (user_id, current_month + '%')).fetchall()

    conn.close()

    total = sum(float(e['amount']) for e in expenses)

    return render_template('report_expense.html',
                           expenses=expenses,
                           cat_totals=cat_totals,
                           total=total,
                           month=datetime.now().strftime('%B %Y'),
                           generated_on=datetime.now().strftime('%d %b %Y, %I:%M %p'))

@app.route('/reports/budget')
def report_budget():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    budgets = conn.execute(
        "SELECT * FROM budget WHERE user_id=?", (user_id,)).fetchall()

    report_data = []
    total_budget = 0
    total_spent = 0
    for b in budgets:
        spent = conn.execute(
            "SELECT SUM(amount) FROM expenses WHERE user_id=? AND category=?",
            (user_id, b['category'])).fetchone()[0] or 0
        spent = float(spent)
        budget_amt = float(b['amount'])
        pct = round((spent / budget_amt) * 100, 1) if budget_amt > 0 else 0
        total_budget += budget_amt
        total_spent += spent
        report_data.append({
            'category': b['category'],
            'month': b['month'],
            'budget': budget_amt,
            'spent': spent,
            'remaining': budget_amt - spent,
            'percentage': pct,
            'status': 'Over Budget' if spent > budget_amt else 'On Track'
        })

    conn.close()

    return render_template('report_budget.html',
                           report_data=report_data,
                           total_budget=total_budget,
                           total_spent=total_spent,
                           generated_on=datetime.now().strftime('%d %b %Y, %I:%M %p'))

@app.route('/reports/investment')
def report_investment():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    inv_list = conn.execute(
        "SELECT * FROM investments WHERE user_id=?", (user_id,)).fetchall()
    conn.close()

    report_data = []
    total_invested = 0
    total_current = 0
    for inv in inv_list:
        invested = float(inv['invested_amount'])
        current = float(inv['current_value'])
        pl = current - invested
        ret = round((pl / invested) * 100, 2) if invested > 0 else 0
        total_invested += invested
        total_current += current
        report_data.append({
            'asset': inv['asset'],
            'name': inv['investment_name'],
            'invested': invested,
            'current': current,
            'pl': pl,
            'return_pct': ret
        })

    total_pl = total_current - total_invested
    overall_roi = round((total_pl / total_invested) * 100, 2) if total_invested > 0 else 0

    return render_template('report_investment.html',
                           report_data=report_data,
                           total_invested=total_invested,
                           total_current=total_current,
                           total_pl=total_pl,
                           overall_roi=overall_roi,
                           generated_on=datetime.now().strftime('%d %b %Y, %I:%M %p'))

@app.route('/reports/goals')
def report_goals():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    goals_list = conn.execute(
        "SELECT * FROM goals WHERE user_id=?", (user_id,)).fetchall()
    conn.close()

    report_data = []
    for g in goals_list:
        target = float(g['target_amount'])
        saved = float(g['saved_amount'])
        pct = round((saved / target) * 100, 1) if target > 0 else 0
        report_data.append({
            'goal_name': g['goal_name'],
            'target': target,
            'saved': saved,
            'remaining': target - saved,
            'percentage': min(pct, 100),
            'target_date': g['target_date'],
            'status': 'Completed' if pct >= 100 else 'In Progress'
        })

    completed_count = sum(1 for g in report_data if g['status'] == 'Completed')

    return render_template('report_goals.html',
                           report_data=report_data,
                           completed_count=completed_count,
                           total_goals=len(report_data),
                           generated_on=datetime.now().strftime('%d %b %Y, %I:%M %p'))

@app.route('/reports/expense/pdf')
def export_expense_pdf():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    current_month = datetime.now().strftime('%Y-%m')

    expenses = conn.execute(
        "SELECT * FROM expenses WHERE user_id=? AND date LIKE ? ORDER BY date DESC",
        (user_id, current_month + '%')).fetchall()

    cat_totals = conn.execute(
        "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? AND date LIKE ? GROUP BY category ORDER BY total DESC",
        (user_id, current_month + '%')).fetchall()

    conn.close()

    total = sum(float(e['amount']) for e in expenses)
    month_name = datetime.now().strftime('%B %Y')

    mem = io.BytesIO()
    doc = SimpleDocTemplate(mem, pagesize=A4,
                            topMargin=0.6*inch, bottomMargin=0.6*inch)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle('TitleStyle', parent=styles['Title'],
                                 textColor=colors.HexColor('#1a1a2e'))
    heading_style = ParagraphStyle('HeadingStyle', parent=styles['Heading2'],
                                   textColor=colors.HexColor('#1a1a2e'),
                                   spaceBefore=16, spaceAfter=8)

    elements = []
    elements.append(Paragraph(f"Monthly Expense Report — {month_name}", title_style))
    elements.append(Paragraph(f"Generated for: {session['user_name']}", styles['Normal']))
    elements.append(Spacer(1, 16))

    elements.append(Paragraph(f"Total Spent: Rs. {total:,.2f}", heading_style))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph("Category Breakdown", heading_style))
    cat_table_data = [['Category', 'Total', '% of Spend']]
    for c in cat_totals:
        pct = round((float(c['total']) / total) * 100, 1) if total > 0 else 0
        cat_table_data.append([c['category'], f"Rs. {float(c['total']):,.2f}", f"{pct}%"])

    cat_table = Table(cat_table_data, colWidths=[2.5*inch, 2*inch, 1.5*inch])
    cat_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a1a2e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
        ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
    ]))
    elements.append(cat_table)
    elements.append(Spacer(1, 16))

    elements.append(Paragraph("All Transactions", heading_style))
    exp_table_data = [['Date', 'Category', 'Amount', 'Description']]
    for e in expenses:
        exp_table_data.append([
            e['date'], e['category'], f"Rs. {float(e['amount']):,.2f}", e['description'] or ''
        ])

    if len(exp_table_data) > 1:
        exp_table = Table(exp_table_data, colWidths=[1.2*inch, 1.5*inch, 1.3*inch, 2*inch])
        exp_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1a1a2e')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#dddddd')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f5f5f5')]),
            ('TOPPADDING', (0, 0), (-1, -1), 5),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ]))
        elements.append(exp_table)
    else:
        elements.append(Paragraph("No expenses recorded this month.", styles['Normal']))

    doc.build(elements)
    mem.seek(0)

    return send_file(mem, mimetype='application/pdf', as_attachment=True,
                     download_name=f'expense_report_{current_month}.pdf')
@app.route('/reports/investment/excel')
def export_investment_excel():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    user_id = session['user_id']
    conn = get_db()

    inv_list = conn.execute(
        "SELECT * FROM investments WHERE user_id=?", (user_id,)).fetchall()
    conn.close()

    wb = Workbook()
    ws = wb.active
    ws.title = "Investment Report"

    ws.append(['Investment Performance Report'])
    ws.append([f"Generated for: {session['user_name']}"])
    ws.append([])
    ws.append(['Asset Type', 'Investment Name', 'Invested Amount', 'Current Value', 'Profit/Loss', 'Return %'])

    total_invested = 0
    total_current = 0

    for inv in inv_list:
        invested = float(inv['invested_amount'])
        current = float(inv['current_value'])
        pl = current - invested
        ret = round((pl / invested) * 100, 2) if invested > 0 else 0
        total_invested += invested
        total_current += current
        ws.append([inv['asset'], inv['investment_name'], invested, current, pl, ret])

    ws.append([])
    total_pl = total_current - total_invested
    overall_roi = round((total_pl / total_invested) * 100, 2) if total_invested > 0 else 0
    ws.append(['TOTAL', '', total_invested, total_current, total_pl, overall_roi])

    mem = io.BytesIO()
    wb.save(mem)
    mem.seek(0)

    return send_file(mem, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name='investment_report.xlsx')
import re

def jarvis_get_response(user_id, message):
    conn = get_db()
    msg = message.lower().strip()

    # Pull common data used across multiple intents
    income_row = conn.execute(
        "SELECT amount FROM income WHERE user_id=? ORDER BY id DESC LIMIT 1",
        (user_id,)).fetchone()
    monthly_income = float(income_row['amount']) if income_row else 0

    total_exp = conn.execute(
        "SELECT SUM(amount) FROM expenses WHERE user_id=?",
        (user_id,)).fetchone()[0] or 0
    total_exp = float(total_exp)

    savings = monthly_income - total_exp
    savings_ratio = round((savings / monthly_income) * 100, 1) if monthly_income > 0 else 0

    # --- Intent: Greeting ---
    if re.search(r'\b(hi|hello|hey|good morning|good evening)\b', msg):
        conn.close()
        return "Hello! I'm JARVIS, your financial assistant. Ask me about your expenses, budget, investments, or goals — I can also give savings tips."

    # --- Intent: Expense summary ---
    if re.search(r'\b(expense|spend|spent|spending)\b', msg):
        cat_expenses = conn.execute(
            "SELECT category, SUM(amount) as total FROM expenses WHERE user_id=? GROUP BY category ORDER BY total DESC",
            (user_id,)).fetchall()
        conn.close()
        if not cat_expenses:
            return "You haven't logged any expenses yet. Add some on the Expenses page and I can summarize them for you."
        top = cat_expenses[0]
        lines = [f"Here's your expense summary: you've spent ₹{total_exp:,.2f} in total."]
        lines.append(f"Your highest spending category is {top['category']} at ₹{float(top['total']):,.2f}.")
        if len(cat_expenses) > 1:
            others = ", ".join(f"{c['category']} (₹{float(c['total']):,.0f})" for c in cat_expenses[1:4])
            lines.append(f"Other categories: {others}.")
        return " ".join(lines)

    # --- Intent: Budget ---
    if re.search(r'\bbudget\b', msg):
        budgets = conn.execute("SELECT * FROM budget WHERE user_id=?", (user_id,)).fetchall()
        conn.close()
        if not budgets:
            return "You haven't set any budgets yet. Head to the Budget page to set category-wise limits, and I can track them for you."
        over = []
        conn2 = get_db()
        for b in budgets:
            spent = conn2.execute(
                "SELECT SUM(amount) FROM expenses WHERE user_id=? AND category=?",
                (user_id, b['category'])).fetchone()[0] or 0
            if float(spent) > float(b['amount']):
                over.append(b['category'])
        conn2.close()
        if over:
            return f"You're over budget in: {', '.join(over)}. Consider cutting back on these categories this month."
        return "Good news — you're within budget in every category right now. Keep it up!"

    # --- Intent: Investments ---
    if re.search(r'\b(invest|investment|portfolio|stocks|mutual fund)\b', msg):
        inv_list = conn.execute("SELECT * FROM investments WHERE user_id=?", (user_id,)).fetchall()
        conn.close()
        if not inv_list:
            return "You don't have any investments logged yet. Add one on the Investments page — I can then analyze your portfolio performance."
        total_invested = sum(float(i['invested_amount']) for i in inv_list)
        total_current = sum(float(i['current_value']) for i in inv_list)
        pl = total_current - total_invested
        roi = round((pl / total_invested) * 100, 2) if total_invested > 0 else 0
        status = "in profit" if pl >= 0 else "at a loss"
        return f"Your portfolio is worth ₹{total_current:,.2f} against ₹{total_invested:,.2f} invested — you're {status} with an overall return of {roi}%."

    # --- Intent: Goals ---
    if re.search(r'\b(goal|target|saving for)\b', msg):
        goals = conn.execute("SELECT * FROM goals WHERE user_id=?", (user_id,)).fetchall()
        conn.close()
        if not goals:
            return "You haven't set any financial goals yet. Head to the Goals page to set one, like an Emergency Fund or Vacation Savings."
        lines = []
        for g in goals:
            target = float(g['target_amount'])
            saved = float(g['saved_amount'])
            pct = round((saved / target) * 100, 1) if target > 0 else 0
            status = "completed" if pct >= 100 else f"{pct}% complete"
            lines.append(f"{g['goal_name']} is {status}")
        return "Here's your goal progress: " + "; ".join(lines) + "."

    # --- Intent: Savings advice ---
    if re.search(r'\b(save more|saving tips|how (can|do) i save|advice|recommend)\b', msg):
        conn.close()
        if savings_ratio >= 30:
            return f"You're already saving {savings_ratio}% of your income — excellent! Consider directing extra savings into investments for long-term growth."
        elif savings_ratio >= 15:
            return f"You're saving {savings_ratio}% of your income. Try the 50/30/20 rule — 50% needs, 30% wants, 20% savings — to push this higher."
        else:
            return f"Your savings rate is only {savings_ratio}%. Start by reviewing your highest spending category and cutting it by 10-15% this month."

    # --- Intent: Health score / overview ---
    if re.search(r'\b(health|score|overview|summary|how am i doing)\b', msg):
        conn.close()
        return (f"Quick overview: Income ₹{monthly_income:,.2f}, Expenses ₹{total_exp:,.2f}, "
                f"Savings ₹{savings:,.2f} ({savings_ratio}% of income). "
                f"{'You are doing well!' if savings_ratio >= 20 else 'There is room to improve your savings rate.'}")

    # --- Default fallback ---
    conn.close()
    return ("I can help with expenses, budget, investments, goals, savings tips, or your overall financial health. "
            "Try asking something like 'How much have I spent?' or 'How are my goals doing?'")


@app.route('/jarvis', methods=['GET'])
def jarvis():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    if 'jarvis_history' not in session:
        session['jarvis_history'] = []
    return render_template('jarvis.html', history=session['jarvis_history'])


@app.route('/jarvis/ask', methods=['POST'])
def jarvis_ask():
    if 'user_id' not in session:
        return {'error': 'not logged in'}, 401

    user_id = session['user_id']
    user_message = request.json.get('message', '').strip()

    if not user_message:
        return {'reply': "Please type a question for me!"}

    reply = jarvis_get_response(user_id, user_message)

    if 'jarvis_history' not in session:
        session['jarvis_history'] = []

    history = session['jarvis_history']
    history.append({'sender': 'user', 'text': user_message})
    history.append({'sender': 'jarvis', 'text': reply})
    session['jarvis_history'] = history[-40:]  # keep last 20 exchanges
    session.modified = True

    return {'reply': reply}


@app.route('/jarvis/clear')
def jarvis_clear():
    if 'user_id' not in session:
        return redirect(url_for('login'))
    session['jarvis_history'] = []
    return redirect(url_for('jarvis'))

init_db()

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port, debug=False)